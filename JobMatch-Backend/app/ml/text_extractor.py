import os
import io
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


class TextExtractor:
    """
    Extracts raw text from PDF, DOCX, and image files.
    Tries multiple methods in order, falls back gracefully.
    """

    SUPPORTED_TYPES = {
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "application/msword": "doc",
        "image/jpeg": "image",
        "image/png": "image",
        "image/jpg": "image",
        "text/plain": "txt",
    }

    # ── Public entry point ───────────────────────────────────────

    def extract(self, file_path: str, mime_type: str = None) -> dict:
        """
        Extract text from a file.
        Returns dict: {success, text, method_used, char_count, error}
        """
        path = Path(file_path)
        if not path.exists():
            return self._result(False, error=f"File not found: {file_path}")

        # Detect type from extension if mime_type not provided
        ext = path.suffix.lower()
        if not mime_type:
            mime_type = self._ext_to_mime(ext)

        file_kind = self.SUPPORTED_TYPES.get(mime_type)

        if file_kind == "pdf":
            return self._extract_pdf(file_path)
        elif file_kind == "docx":
            return self._extract_docx(file_path)
        elif file_kind == "image":
            return self._extract_image(file_path)
        elif file_kind == "txt":
            return self._extract_txt(file_path)
        else:
            return self._result(False, error=f"Unsupported file type: {mime_type}")

    # ── PDF extraction ───────────────────────────────────────────

    def _extract_pdf(self, file_path: str) -> dict:
        # Method 1: pdfplumber (best for structured PDFs)
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text.strip())
            text = "\n\n".join(text_parts)
            if len(text.strip()) > 100:
                return self._result(True, text=text, method="pdfplumber")
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}")

        # Method 2: PyPDF2 (fallback)
        try:
            import PyPDF2
            text_parts = []
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text.strip())
            text = "\n\n".join(text_parts)
            if len(text.strip()) > 100:
                return self._result(True, text=text, method="PyPDF2")
        except Exception as e:
            logger.warning(f"PyPDF2 failed: {e}")

        # Method 3: OCR via pytesseract (scanned PDFs)
        try:
            return self._pdf_ocr(file_path)
        except Exception as e:
            logger.warning(f"PDF OCR failed: {e}")

        return self._result(False, error="All PDF extraction methods failed")

    def _pdf_ocr(self, file_path: str) -> dict:
        """Convert PDF pages to images and OCR each one."""
        try:
            from pdf2image import convert_from_path
            import pytesseract

            images = convert_from_path(file_path, dpi=200)
            text_parts = []
            for img in images:
                page_text = pytesseract.image_to_string(img, lang="eng")
                if page_text.strip():
                    text_parts.append(page_text.strip())
            text = "\n\n".join(text_parts)
            return self._result(bool(text.strip()), text=text, method="pdf_ocr")
        except ImportError:
            # pdf2image not installed — skip silently
            return self._result(False, error="pdf2image not available for OCR")

    # ── DOCX extraction ──────────────────────────────────────────

    def _extract_docx(self, file_path: str) -> dict:
        try:
            from docx import Document
            doc = Document(file_path)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

            # Also extract table content
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())

            text = "\n".join(paragraphs)
            return self._result(bool(text.strip()), text=text, method="python-docx")
        except Exception as e:
            return self._result(False, error=f"DOCX extraction failed: {e}")

    # ── Image extraction (OCR) ───────────────────────────────────

    def _extract_image(self, file_path: str) -> dict:
        try:
            import pytesseract
            from PIL import Image, ImageFilter, ImageEnhance

            img = Image.open(file_path)

            # Pre-process for better OCR
            img = img.convert("L")                          # Grayscale
            img = img.filter(ImageFilter.SHARPEN)           # Sharpen
            enhancer = ImageEnhance.Contrast(img)
            img = enhancer.enhance(2.0)                     # Boost contrast

            text = pytesseract.image_to_string(img, lang="eng")
            return self._result(bool(text.strip()), text=text.strip(), method="pytesseract")
        except Exception as e:
            return self._result(False, error=f"Image OCR failed: {e}")

    # ── Plain text ───────────────────────────────────────────────

    def _extract_txt(self, file_path: str) -> dict:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            return self._result(bool(text.strip()), text=text.strip(), method="plaintext")
        except Exception as e:
            return self._result(False, error=f"TXT read failed: {e}")

    # ── Helpers ──────────────────────────────────────────────────

    def _result(self, success: bool, text: str = None, method: str = "none", error: str = None) -> dict:
        text = self._clean_text(text) if text else ""
        return {
            "success": success,
            "text": text,
            "method_used": method,
            "char_count": len(text),
            "error": error,
        }

    def _clean_text(self, text: str) -> str:
        """Normalize whitespace and remove garbage characters."""
        if not text:
            return ""
        import re
        # Remove non-printable characters
        text = re.sub(r"[^\x20-\x7E\n\r\t]", " ", text)
        # Collapse multiple spaces
        text = re.sub(r" {3,}", "  ", text)
        # Collapse more than 3 consecutive newlines
        text = re.sub(r"\n{4,}", "\n\n\n", text)
        return text.strip()

    def _ext_to_mime(self, ext: str) -> str:
        mapping = {
            ".pdf":  "application/pdf",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".doc":  "application/msword",
            ".png":  "image/png",
            ".jpg":  "image/jpeg",
            ".jpeg": "image/jpeg",
            ".txt":  "text/plain",
        }
        return mapping.get(ext, "application/octet-stream")


# Singleton instance
text_extractor = TextExtractor()